import docx


def extract_text_docx(file_path):
    """
    Extrai o texto de um arquivo DOCX.

    :param file_path: Caminho para o arquivo DOCX
    :return: Texto extraído do documento
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Extrair texto das tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return '\n'.join(full_text)
    except Exception as e:
        print(f"Erro ao extrair texto do arquivo DOCX {file_path}: {str(e)}")
        return ""


# Exemplo de uso
if __name__ == "__main__":
    docx_path = "caminho/para/seu/arquivo.docx"
    extracted_text = extract_text_docx(docx_path)
    print(extracted_text)